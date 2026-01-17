
import os
import re
from dataclasses import dataclass
from pathlib import Path

from app.core.config import settings
from app.core.exceptions import DocumentParsingError, UnsupportedFileTypeError
from app.core.logging import get_logger

logger = get_logger(__name__)


@dataclass
class ExtractedDocument:
    
    text: str
    filename: str
    file_type: str
    page_count: int | None = None
    title: str | None = None
    headings: list[str] | None = None
    extraction_warnings: list[str] | None = None


class DocumentProcessor:
    
    def __init__(self) -> None:
        self.supported_extensions = settings.supported_extensions_list
    
    def process_file(self, file_path: str | Path) -> ExtractedDocument:
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise DocumentParsingError(
                filename=str(file_path),
                reason="File does not exist",
            )
        
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_extensions:
            raise UnsupportedFileTypeError(
                file_type=extension,
                supported_types=self.supported_extensions,
            )
        
        logger.info(
            "Processing document",
            filename=file_path.name,
            file_type=extension,
        )
        
        # Route to appropriate handler
        if extension == ".pdf":
            return self._process_pdf(file_path)
        elif extension == ".docx":
            return self._process_docx(file_path)
        elif extension == ".md":
            return self._process_markdown(file_path)
        elif extension == ".txt":
            return self._process_text(file_path)
        else:
            raise UnsupportedFileTypeError(
                file_type=extension,
                supported_types=self.supported_extensions,
            )
    
    def _process_pdf(self, file_path: Path) -> ExtractedDocument:
        try:
            from pypdf import PdfReader
        except ImportError:
            raise DocumentParsingError(
                filename=file_path.name,
                reason="pypdf library not installed",
            )
        
        warnings = []
        
        try:
            reader = PdfReader(str(file_path))
            
            # Extract text from all pages
            pages_text = []
            for i, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        pages_text.append(text)
                    else:
                        warnings.append(f"Page {i+1} had no extractable text")
                except Exception as e:
                    warnings.append(f"Error extracting page {i+1}: {str(e)}")
            
            full_text = "\n\n".join(pages_text)
            
            # Try to extract title from metadata
            title = None
            if reader.metadata:
                title = reader.metadata.get("/Title")
            
            # Clean up the text
            full_text = self._clean_text(full_text)
            
            logger.info(
                "PDF processed successfully",
                filename=file_path.name,
                pages=len(reader.pages),
                text_length=len(full_text),
            )
            
            return ExtractedDocument(
                text=full_text,
                filename=file_path.name,
                file_type="pdf",
                page_count=len(reader.pages),
                title=title,
                extraction_warnings=warnings if warnings else None,
            )
            
        except Exception as e:
            logger.error(
                "PDF extraction failed",
                filename=file_path.name,
                error=str(e),
            )
            raise DocumentParsingError(
                filename=file_path.name,
                reason=f"PDF extraction failed: {str(e)}",
            )
    
    def _process_docx(self, file_path: Path) -> ExtractedDocument:
        try:
            from docx import Document
        except ImportError:
            raise DocumentParsingError(
                filename=file_path.name,
                reason="python-docx library not installed",
            )
        
        try:
            doc = Document(str(file_path))
            
            # Extract text from paragraphs
            paragraphs = []
            headings = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
                    
                    # Track headings
                    if para.style and para.style.name.startswith("Heading"):
                        headings.append(text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        paragraphs.append(row_text)
            
            full_text = "\n\n".join(paragraphs)
            full_text = self._clean_text(full_text)
            
            # Try to extract title from core properties
            title = None
            if doc.core_properties.title:
                title = doc.core_properties.title
            elif headings:
                title = headings[0]  # Use first heading as title
            
            logger.info(
                "DOCX processed successfully",
                filename=file_path.name,
                paragraphs=len(paragraphs),
                text_length=len(full_text),
            )
            
            return ExtractedDocument(
                text=full_text,
                filename=file_path.name,
                file_type="docx",
                title=title,
                headings=headings if headings else None,
            )
            
        except Exception as e:
            logger.error(
                "DOCX extraction failed",
                filename=file_path.name,
                error=str(e),
            )
            raise DocumentParsingError(
                filename=file_path.name,
                reason=f"DOCX extraction failed: {str(e)}",
            )
    
    def _process_markdown(self, file_path: Path) -> ExtractedDocument:
        try:
            # Read raw markdown
            with open(file_path, "r", encoding="utf-8") as f:
                raw_text = f.read()
            
            # Extract headings
            headings = re.findall(r"^#+\s+(.+)$", raw_text, re.MULTILINE)
            
            # Convert markdown to plain text (remove formatting)
            text = self._markdown_to_text(raw_text)
            text = self._clean_text(text)
            
            # Use first heading as title
            title = headings[0] if headings else None
            
            logger.info(
                "Markdown processed successfully",
                filename=file_path.name,
                text_length=len(text),
            )
            
            return ExtractedDocument(
                text=text,
                filename=file_path.name,
                file_type="md",
                title=title,
                headings=headings if headings else None,
            )
            
        except Exception as e:
            logger.error(
                "Markdown extraction failed",
                filename=file_path.name,
                error=str(e),
            )
            raise DocumentParsingError(
                filename=file_path.name,
                reason=f"Markdown extraction failed: {str(e)}",
            )
    
    def _process_text(self, file_path: Path) -> ExtractedDocument:
        try:
            # Try UTF-8 first, fall back to other encodings
            encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
            
            text = None
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        text = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if text is None:
                raise DocumentParsingError(
                    filename=file_path.name,
                    reason="Could not decode file with any supported encoding",
                )
            
            text = self._clean_text(text)
            
            logger.info(
                "Text file processed successfully",
                filename=file_path.name,
                text_length=len(text),
            )
            
            return ExtractedDocument(
                text=text,
                filename=file_path.name,
                file_type="txt",
            )
            
        except Exception as e:
            if isinstance(e, DocumentParsingError):
                raise
            logger.error(
                "Text extraction failed",
                filename=file_path.name,
                error=str(e),
            )
            raise DocumentParsingError(
                filename=file_path.name,
                reason=f"Text extraction failed: {str(e)}",
            )
    
    def _markdown_to_text(self, markdown: str) -> str:
        # Remove code blocks
        text = re.sub(r"```[\s\S]*?```", "", markdown)
        text = re.sub(r"`[^`]+`", "", text)
        
        # Convert headers to plain text
        text = re.sub(r"^#+\s+", "", text, flags=re.MULTILINE)
        
        # Remove emphasis markers
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"\*([^*]+)\*", r"\1", text)
        text = re.sub(r"__([^_]+)__", r"\1", text)
        text = re.sub(r"_([^_]+)_", r"\1", text)
        
        # Remove links, keep text
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
        
        # Remove images
        text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", "", text)
        
        # Remove horizontal rules
        text = re.sub(r"^[-*_]{3,}$", "", text, flags=re.MULTILINE)
        
        return text
    
    def _clean_text(self, text: str) -> str:
        # Normalize line endings
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        
        # Replace tabs with spaces
        text = text.replace("\t", " ")
        
        # Remove excessive whitespace within lines
        text = re.sub(r" +", " ", text)
        
        # Remove excessive blank lines (keep max 2)
        text = re.sub(r"\n{3,}", "\n\n", text)
        
        # Strip lines
        lines = [line.strip() for line in text.split("\n")]
        text = "\n".join(lines)
        
        # Final strip
        text = text.strip()
        
        return text
    
    def process_directory(
        self,
        directory: str | Path,
        recursive: bool = True,
    ) -> list[ExtractedDocument]:
        directory = Path(directory)
        
        if not directory.is_dir():
            raise DocumentParsingError(
                filename=str(directory),
                reason="Path is not a directory",
            )
        
        results = []
        errors = []
        
        # Find all supported files
        pattern = "**/*" if recursive else "*"
        for ext in self.supported_extensions:
            for file_path in directory.glob(f"{pattern}{ext}"):
                try:
                    result = self.process_file(file_path)
                    results.append(result)
                except Exception as e:
                    errors.append(f"{file_path}: {str(e)}")
                    logger.warning(
                        "Failed to process file",
                        file_path=str(file_path),
                        error=str(e),
                    )
        
        logger.info(
            "Directory processing complete",
            directory=str(directory),
            processed=len(results),
            errors=len(errors),
        )
        
        return results
